# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def create_document(name, id_number, date):
    # Create a new Document
    doc = Document()

    # Add a bold and centered heading
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading.add_run("Výsledný protokol genetického vyšetření")
    heading_run.bold = True
    heading_run.font.size = Pt(16)

    # Add a table with the provided information
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # Fill in the table with information
    table.cell(0, 0).text = "Jméno a příjmení:"
    table.cell(0, 1).text = name

    table.cell(1, 0).text = "Rodné číslo:"
    table.cell(1, 1).text = id_number

    table.cell(2, 0).text = "Datum odběru:"
    table.cell(2, 1).text = date

    # Save the document
    doc.save(r'C:\Users\irina\OneDrive\Documents\Programming\Prakticky ukol\document_ukol2.docx')

if __name__ == "__main__":
    # Example usage:
    name_argument = "argument 1"
    id_number_argument = "argument 2"
    date_argument = "argument 3"

    create_document(name_argument, id_number_argument, date_argument)
